import io
import docx
import pypdf

class ResumeParsingError(Exception):
    """Custom exception raised when resume parsing fails."""
    pass

def extract_pdf_text(file_bytes: bytes) -> str:
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = pypdf.PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    except Exception as e:
        raise ResumeParsingError(f"Failed to parse PDF resume: {str(e)}")

def extract_docx_text(file_bytes: bytes) -> str:
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = docx.Document(docx_file)
        paragraphs_text = [p.text for p in doc.paragraphs]

        # Extract text from tables if any (resumes frequently use table structures)
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in table_text:
                        table_text.append(cell_text)

        full_text = "\n".join(paragraphs_text + table_text)
        return full_text
    except Exception as e:
        raise ResumeParsingError(f"Failed to parse DOCX resume: {str(e)}")

def extract_txt_text(file_bytes: bytes) -> str:
    encodings = ["utf-8", "latin-1", "utf-16", "utf-16le", "utf-16be"]
    for encoding in encodings:
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ResumeParsingError("Failed to decode TXT resume. Unsupported file encoding.")

def extract_resume_text(filename: str, file_bytes: bytes) -> str:
    if not file_bytes or len(file_bytes.strip()) == 0:
        raise ResumeParsingError("The uploaded resume file is empty.")

    ext = filename.lower().split(".")[-1]
    if ext == "pdf":
        raw_text = extract_pdf_text(file_bytes)
    elif ext == "docx":
        raw_text = extract_docx_text(file_bytes)
    elif ext in ["txt", "text"]:
        raw_text = extract_txt_text(file_bytes)
    else:
        raise ResumeParsingError(f"Unsupported file format: .{ext}")

    # Remove repeated blank lines and trim whitespace
    lines = [line.strip() for line in raw_text.splitlines()]
    cleaned_lines = []
    for line in lines:
        if line:
            cleaned_lines.append(line)
        else:
            if cleaned_lines and cleaned_lines[-1] != "":
                cleaned_lines.append("")

    cleaned_text = "\n".join(cleaned_lines).strip()
    if not cleaned_text:
        raise ResumeParsingError("No readable text content found in the resume.")

    return cleaned_text
